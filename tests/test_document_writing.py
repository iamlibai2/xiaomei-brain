import base64
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor

from xiaomei_brain.plugin.context import PluginContext
from xiaomei_brain.plugin.loader import PluginLoader
from xiaomei_brain.plugin.registry import PluginRegistry
from xiaomei_brain.plugins.tools.document_word.adapter import register as register_word
from xiaomei_brain.tools.builtin.documents import create_write_document_tool
from xiaomei_brain.tools.execution_context import bind_tool_execution
from xiaomei_brain.gateway import artifacts as artifact_module
from xiaomei_brain.gateway.artifacts import discover_tool_artifacts


def _word_registry() -> PluginRegistry:
    registry = PluginRegistry()
    context = PluginContext({}, "document_word", "test", registry)
    register_word(context)
    return registry


def test_word_plugin_owns_writer_and_skill_directory():
    registry = _word_registry()

    assert registry.get_document_writer("word") is not None
    assert registry.list_document_writers() == ["word"]
    skill_dirs = registry.get_skill_directories()
    assert len(skill_dirs) == 1
    assert (Path(skill_dirs[0]) / "SKILL.md").is_file()


def test_word_theme_preview_is_copied_into_current_execution_output(tmp_path):
    registry = _word_registry()
    tool = next(
        item
        for item in registry.get_agent_tools()
        if item.name == "preview_word_themes"
    )
    workspace = tmp_path / "workspace"
    output_root = workspace / "work"
    output_root.mkdir(parents=True)

    with bind_tool_execution(
        tool_call_id="call-preview",
        tool_name=tool.name,
        arguments={},
        artifact_callback=None,
        session_id="desktop-session-1",
        workspace_root=str(workspace),
        working_directory=str(output_root),
        output_root=str(output_root),
    ):
        result = tool.execute()

    assert result["success"] is True
    assert result["themes"] == [
        "business-blue",
        "modern-minimal",
        "warm-professional",
        "technology",
    ]
    preview = Path(result["output_path"])
    assert preview.name == "Word主题预览.png"
    assert preview.resolve().is_relative_to(output_root.resolve())
    assert preview.read_bytes().startswith(b"\x89PNG\r\n\x1a\n")


def test_bundled_word_theme_previews_are_real_png_assets():
    assets = (
        Path(__file__).parents[1]
        / "src"
        / "xiaomei_brain"
        / "plugins"
        / "tools"
        / "document_word"
        / "assets"
    )
    expected = {
        "business-blue.png",
        "modern-minimal.png",
        "warm-professional.png",
        "technology.png",
        "theme-showcase.png",
    }
    assert {path.name for path in assets.glob("*.png")} == expected
    assert all((assets / name).stat().st_size > 10_000 for name in expected)


def test_document_tools_and_writer_are_discovered_through_plugins():
    registry = PluginRegistry()
    tools_root = (
        Path(__file__).parents[1]
        / "src"
        / "xiaomei_brain"
        / "plugins"
        / "tools"
    )
    loader = PluginLoader(
        registry,
        config={
            "plugins": {
                "allow": [
                    "document_io",
                    "document_pdf",
                    "document_presentation",
                    "document_spreadsheet",
                    "document_word",
                ],
            },
        },
        agent_id="test",
    )

    loaded = loader.boot([str(tools_root)])

    assert {
        item.manifest.name
        for item in loaded
        if item.status == "loaded"
    } == {
        "document_io",
        "document_pdf",
        "document_presentation",
        "document_spreadsheet",
        "document_word",
    }
    assert {
        tool.name for tool in registry.get_agent_tools()
    } == {
        "read_document",
        "write_document",
        "preview_word_themes",
        "manage_document_template",
    }
    word_writer = registry.get_document_writer("word")
    assert word_writer is not None
    assert word_writer.template_analyzer is not None
    assert registry.get_document_writer("spreadsheet") is not None
    assert registry.get_document_writer("presentation") is not None
    assert registry.get_document_writer("pdf") is not None


def test_write_document_creates_and_validates_word_file(tmp_path):
    registry = _word_registry()
    tool = create_write_document_tool(registry)
    workspace = tmp_path / "workspace"
    outputs = workspace / "outputs"
    workspace.mkdir()
    spec = workspace / "report.json"
    spec.write_text(json.dumps({
        "title": "Quarterly Report",
        "properties": {"author": "Xiaomei"},
        "blocks": [
            {"type": "heading", "level": 1, "text": "Summary"},
            {"type": "paragraph", "text": "Work completed."},
            {"type": "list", "items": ["One", "Two"]},
            {"type": "table", "headers": ["Metric", "Value"], "rows": [["Users", 12]]},
        ],
    }), encoding="utf-8")

    with bind_tool_execution(
        tool_call_id="call-1",
        tool_name="write_document",
        arguments={},
        artifact_callback=None,
        workspace_root=str(workspace),
        output_root=str(outputs),
    ):
        result = tool.execute(
            format="word",
            specification_path="report.json",
            output_name="quarterly.docx",
        )

    output = outputs / "quarterly.docx"
    assert result["success"] is True
    assert result["validation"]["valid"] is True
    assert result["validation"]["tables"] == 1
    assert output.is_file()
    document = Document(output)
    assert document.core_properties.author == "Xiaomei"
    assert "Quarterly Report" in "\n".join(p.text for p in document.paragraphs)


def test_word_creation_applies_professional_theme_and_table_hierarchy(tmp_path):
    registry = _word_registry()
    tool = create_write_document_tool(registry)
    workspace = tmp_path / "workspace"
    outputs = workspace / "outputs"
    workspace.mkdir()
    spec = workspace / "themed.json"
    spec.write_text(json.dumps({
        "title": "企业试点报告",
        "subtitle": "管理层决策材料",
        "theme": {"preset": "technology"},
        "blocks": [
            {"type": "heading", "level": 1, "text": "执行摘要"},
            {"type": "paragraph", "text": "本报告用于验证专业排版主题。"},
            {"type": "quote", "text": "结论先行，证据随后。"},
            {
                "type": "table",
                "headers": ["指标", "结果"],
                "rows": [["效率", "提升20%"], ["风险", "可控"]],
                "column_widths_cm": [5, 8],
            },
        ],
    }, ensure_ascii=False), encoding="utf-8")

    with bind_tool_execution(
        tool_call_id="call-theme",
        tool_name="write_document",
        arguments={},
        artifact_callback=None,
        workspace_root=str(workspace),
        output_root=str(outputs),
    ):
        result = tool.execute(
            format="word",
            specification_path="themed.json",
            output_name="themed.docx",
        )

    assert result["success"] is True
    assert result["validation"]["theme"] == "technology"
    assert result["validation"]["render_validation"]["status"] == "disabled"
    document = Document(outputs / "themed.docx")
    assert round(document.sections[0].page_width.cm, 1) == 21.0
    assert round(document.sections[0].page_height.cm, 1) == 29.7
    assert document.styles["Title"].font.color.rgb == RGBColor(0x12, 0x3B, 0x5D)
    assert document.styles["Heading 1"].paragraph_format.keep_with_next is True
    title_fonts = document.styles["Title"]._element.rPr.rFonts
    assert title_fonts.get(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
    ) == "Microsoft YaHei"
    assert "eastAsiaTheme" not in title_fonts.xml
    assert "w:pBdr" in document.styles["Heading 1"]._element.xml
    assert "w:fill=\"123B5D\"" in document.tables[0].rows[0].cells[0]._tc.xml
    assert document.tables[0].rows[0].cells[0].paragraphs[0].runs[0].bold is True


def test_word_visual_validation_is_opt_in_and_reported(tmp_path, monkeypatch):
    registry = _word_registry()
    tool = create_write_document_tool(registry)
    workspace = tmp_path / "workspace"
    outputs = workspace / "outputs"
    workspace.mkdir()
    spec = workspace / "visual.json"
    spec.write_text(json.dumps({
        "visual_validation": True,
        "blocks": [{"type": "paragraph", "text": "需要渲染检查"}],
    }, ensure_ascii=False), encoding="utf-8")

    monkeypatch.setattr(
        "xiaomei_brain.plugins.tools.document_word.writer.render_office_document",
        lambda path: {
            "status": "passed",
            "performed": True,
            "backend": "microsoft-office-com",
            "page_count": 1,
            "blank_pages": [],
        },
    )

    with bind_tool_execution(
        tool_call_id="call-visual",
        tool_name="write_document",
        arguments={},
        artifact_callback=None,
        workspace_root=str(workspace),
        output_root=str(outputs),
    ):
        result = tool.execute(
            format="word",
            specification_path="visual.json",
            output_name="visual.docx",
        )

    rendered = result["validation"]["render_validation"]
    assert rendered["status"] == "passed"
    assert rendered["backend"] == "microsoft-office-com"


def test_write_document_revises_copy_without_overwriting_attachment(tmp_path):
    registry = _word_registry()
    tool = create_write_document_tool(registry)
    workspace = tmp_path / "workspace"
    outputs = workspace / "outputs"
    workspace.mkdir()
    source = tmp_path / "source.docx"
    original = Document()
    original.add_paragraph("Old wording")
    original.save(source)
    original_bytes = source.read_bytes()
    spec = workspace / "revision.json"
    spec.write_text(json.dumps({
        "operations": [
            {"type": "replace_text", "old": "Old wording", "new": "New wording"},
            {"type": "append_blocks", "blocks": [{"type": "paragraph", "text": "Added note"}]},
        ],
    }), encoding="utf-8")
    attachment = {
        "id": "source-1",
        "name": "source.docx",
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "size": source.stat().st_size,
        "kind": "document",
        "local_path": str(source),
    }

    with bind_tool_execution(
        tool_call_id="call-2",
        tool_name="write_document",
        arguments={},
        artifact_callback=None,
        attachments=(attachment,),
        workspace_root=str(workspace),
        output_root=str(outputs),
    ):
        result = tool.execute(
            format="word",
            specification_path="revision.json",
            output_name="revised.docx",
            source_attachment_id="source-1",
        )

    assert result["success"] is True
    assert result["validation"]["replacements"] == 1
    assert source.read_bytes() == original_bytes
    revised = Document(outputs / "revised.docx")
    text = "\n".join(paragraph.text for paragraph in revised.paragraphs)
    assert "New wording" in text and "Added note" in text


def test_write_document_styles_only_selected_existing_table_cells(tmp_path):
    registry = _word_registry()
    tool = create_write_document_tool(registry)
    workspace = tmp_path / "workspace"
    outputs = workspace / "outputs"
    workspace.mkdir()
    source = tmp_path / "source.docx"
    original = Document()
    target = original.add_table(rows=2, cols=2)
    target.cell(0, 0).text = "Region"
    target.cell(0, 1).text = "Revenue"
    target.cell(1, 0).text = "East"
    target.cell(1, 1).text = "100"
    untouched = original.add_table(rows=1, cols=1)
    untouched.cell(0, 0).text = "Keep me"
    original.save(source)
    original_bytes = source.read_bytes()
    spec = workspace / "table-style.json"
    spec.write_text(json.dumps({
        "operations": [{
            "type": "style_table_cells",
            "table_index": 1,
            "rows": [1],
            "columns": "all",
            "fill_color": "D9E2F3",
            "text_color": "1F2937",
            "bold": True,
            "horizontal_alignment": "center",
            "vertical_alignment": "center",
        }],
    }), encoding="utf-8")
    attachment = {
        "id": "source-table",
        "name": "source.docx",
        "kind": "document",
        "local_path": str(source),
    }

    with bind_tool_execution(
        tool_call_id="call-table-style",
        tool_name="write_document",
        arguments={},
        artifact_callback=None,
        attachments=(attachment,),
        workspace_root=str(workspace),
        output_root=str(outputs),
    ):
        result = tool.execute(
            format="word",
            specification_path="table-style.json",
            output_name="styled.docx",
            source_attachment_id="source-table",
        )

    assert result["success"] is True
    assert result["validation"]["styled_cells"] == 2
    assert source.read_bytes() == original_bytes
    revised = Document(outputs / "styled.docx")
    for cell in revised.tables[0].rows[0].cells:
        shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
        assert shading is not None and shading.get(qn("w:fill")) == "D9E2F3"
        run = cell.paragraphs[0].runs[0]
        assert run.bold is True
        assert run.font.color.rgb == RGBColor(0x1F, 0x29, 0x37)
        assert cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert revised.tables[0].cell(1, 0)._tc.get_or_add_tcPr().find(qn("w:shd")) is None
    assert revised.tables[1].cell(0, 0)._tc.get_or_add_tcPr().find(qn("w:shd")) is None


def test_write_document_rejects_out_of_range_table_style_target(tmp_path):
    registry = _word_registry()
    tool = create_write_document_tool(registry)
    workspace = tmp_path / "workspace"
    outputs = workspace / "outputs"
    workspace.mkdir()
    source = tmp_path / "source.docx"
    document = Document()
    document.add_table(rows=1, cols=1).cell(0, 0).text = "Only table"
    document.save(source)
    (workspace / "invalid-table-style.json").write_text(json.dumps({
        "operations": [{
            "type": "style_table_cells",
            "table_index": 2,
            "rows": [1],
            "fill_color": "D9E2F3",
        }],
    }), encoding="utf-8")

    with bind_tool_execution(
        tool_call_id="call-invalid-table-style",
        tool_name="write_document",
        arguments={},
        artifact_callback=None,
        attachments=({
            "id": "source-table",
            "name": "source.docx",
            "kind": "document",
            "local_path": str(source),
        },),
        workspace_root=str(workspace),
        output_root=str(outputs),
    ):
        result = tool.execute(
            format="word",
            specification_path="invalid-table-style.json",
            output_name="styled.docx",
            source_attachment_id="source-table",
        )

    assert "当前文档共有 1 张表格" in result["error"]
    assert not (outputs / "styled.docx").exists()


def test_word_template_replaces_cross_run_placeholders_across_document_parts(tmp_path):
    registry = _word_registry()
    tool = create_write_document_tool(registry)
    workspace = tmp_path / "workspace"
    outputs = workspace / "outputs"
    workspace.mkdir()
    source = tmp_path / "template.docx"

    template = Document()
    paragraph = template.add_paragraph()
    paragraph.add_run("客户：")
    paragraph.add_run("{{customer").bold = True
    paragraph.add_run("_name}}").italic = True
    paragraph.add_run("（重点客户）").underline = True
    table = template.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "项目：{{project_name}}"
    header = template.sections[0].header.paragraphs[0]
    header.text = "报告：{{project_name}}"
    footer = template.sections[0].footer.paragraphs[0]
    footer.text = "联系人：{{customer_name}}"
    template.save(source)

    spec = workspace / "template-values.json"
    spec.write_text(json.dumps({
        "operations": [{
            "type": "replace_placeholders",
            "values": {
                "customer_name": "星海科技",
                "project_name": "智能办公平台",
            },
        }],
    }, ensure_ascii=False), encoding="utf-8")
    attachment = {
        "id": "template-1",
        "name": "template.docx",
        "kind": "document",
        "local_path": str(source),
    }

    with bind_tool_execution(
        tool_call_id="call-template",
        tool_name="write_document",
        arguments={},
        artifact_callback=None,
        attachments=(attachment,),
        workspace_root=str(workspace),
        output_root=str(outputs),
    ):
        result = tool.execute(
            format="word",
            specification_path="template-values.json",
            output_name="filled.docx",
            source_attachment_id="template-1",
        )

    assert result["success"] is True
    assert result["validation"]["replacements"] == 4
    filled = Document(outputs / "filled.docx")
    body = filled.paragraphs[0]
    assert body.text == "客户：星海科技（重点客户）"
    assert body.runs[1].bold is True
    assert body.runs[3].underline is True
    assert filled.tables[0].cell(0, 0).text == "项目：智能办公平台"
    assert filled.sections[0].header.paragraphs[0].text == "报告：智能办公平台"
    assert filled.sections[0].footer.paragraphs[0].text == "联系人：星海科技"


def test_word_template_inserts_blocks_at_body_marker(tmp_path):
    registry = _word_registry()
    tool = create_write_document_tool(registry)
    workspace = tmp_path / "workspace"
    outputs = workspace / "outputs"
    workspace.mkdir()
    source = tmp_path / "template.docx"

    template = Document()
    template.add_paragraph("前言")
    template.add_paragraph("{{DETAILS}}")
    template.add_paragraph("结语")
    template.save(source)
    spec = workspace / "insert.json"
    spec.write_text(json.dumps({
        "operations": [{
            "type": "insert_blocks_after",
            "marker": "{{DETAILS}}",
            "remove_marker": True,
            "blocks": [
                {"type": "heading", "level": 1, "text": "项目详情"},
                {"type": "paragraph", "text": "这是插入的正文。"},
                {"type": "list", "items": ["第一项", "第二项"]},
            ],
        }],
    }, ensure_ascii=False), encoding="utf-8")
    attachment = {
        "id": "template-2",
        "name": "template.docx",
        "kind": "document",
        "local_path": str(source),
    }

    with bind_tool_execution(
        tool_call_id="call-insert",
        tool_name="write_document",
        arguments={},
        artifact_callback=None,
        attachments=(attachment,),
        workspace_root=str(workspace),
        output_root=str(outputs),
    ):
        result = tool.execute(
            format="word",
            specification_path="insert.json",
            output_name="inserted.docx",
            source_attachment_id="template-2",
        )

    assert result["success"] is True
    assert result["validation"]["insertions"] == 4
    inserted = Document(outputs / "inserted.docx")
    assert [paragraph.text for paragraph in inserted.paragraphs] == [
        "前言",
        "项目详情",
        "这是插入的正文。",
        "第一项",
        "第二项",
        "结语",
    ]


def test_word_creation_uses_owned_image_and_page_layout(tmp_path):
    registry = _word_registry()
    tool = create_write_document_tool(registry)
    workspace = tmp_path / "workspace"
    outputs = workspace / "outputs"
    workspace.mkdir()
    image = tmp_path / "logo.png"
    image.write_bytes(base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwC"
        "AAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
    ))
    spec = workspace / "styled.json"
    spec.write_text(json.dumps({
        "default_style": {"font": "Microsoft YaHei", "size_pt": 11},
        "page": {
            "size": "A4",
            "orientation": "landscape",
            "margins_cm": {"top": 2, "right": 1.5, "bottom": 2, "left": 1.5},
        },
        "header": {"text": "小美企业报告"},
        "footer": {"text": "第 ", "page_number": True},
        "blocks": [
            {"type": "paragraph", "text": "报告正文"},
            {
                "type": "image",
                "attachment_id": "logo-1",
                "width_cm": 2,
                "caption": "企业标识",
            },
        ],
    }, ensure_ascii=False), encoding="utf-8")
    attachment = {
        "id": "logo-1",
        "name": "logo.png",
        "kind": "image",
        "local_path": str(image),
    }

    with bind_tool_execution(
        tool_call_id="call-layout",
        tool_name="write_document",
        arguments={},
        artifact_callback=None,
        attachments=(attachment,),
        workspace_root=str(workspace),
        output_root=str(outputs),
    ):
        result = tool.execute(
            format="word",
            specification_path="styled.json",
            output_name="styled.docx",
        )

    assert result["success"] is True
    styled = Document(outputs / "styled.docx")
    section = styled.sections[0]
    assert section.orientation == WD_ORIENT.LANDSCAPE
    assert round(section.top_margin.cm, 1) == 2.0
    assert round(section.left_margin.cm, 1) == 1.5
    assert styled.styles["Normal"].font.name == "Microsoft YaHei"
    assert round(styled.styles["Normal"].font.size.pt, 1) == 11.0
    assert section.header.paragraphs[0].text == "小美企业报告"
    assert section.footer.paragraphs[0].text == "第 "
    footer_xml = section.footer.paragraphs[0]._p.xml
    assert "PAGE" in footer_xml
    assert len(styled.inline_shapes) == 1
    assert any("企业标识" in paragraph.text for paragraph in styled.paragraphs)


def test_word_creation_uses_controlled_workspace_image(tmp_path):
    registry = _word_registry()
    tool = create_write_document_tool(registry)
    workspace = tmp_path / "workspace"
    work = workspace / "work"
    outputs = workspace / "outputs"
    work.mkdir(parents=True)
    image = work / "generated.png"
    image.write_bytes(base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwC"
        "AAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
    ))
    spec = work / "with-image.json"
    spec.write_text(json.dumps({
        "blocks": [{
            "type": "image",
            "workspace_path": "work/generated.png",
            "width_cm": 2,
        }],
    }), encoding="utf-8")

    with bind_tool_execution(
        tool_call_id="call-workspace-image",
        tool_name="write_document",
        arguments={},
        artifact_callback=None,
        workspace_root=str(workspace),
        working_directory=str(work),
        output_root=str(outputs),
    ):
        result = tool.execute(
            format="word",
            specification_path="work/with-image.json",
            output_name="workspace-image.docx",
        )

    assert result["success"] is True
    assert len(Document(outputs / "workspace-image.docx").inline_shapes) == 1


def test_word_long_document_supports_toc_numbering_sections_and_captions(
    tmp_path,
    monkeypatch,
):
    from docx.oxml.ns import qn

    monkeypatch.setattr(
        "xiaomei_brain.plugins.tools.document_word.writer.refresh_word_fields",
        lambda path: {
            "status": "updated",
            "performed": True,
            "backend": "test-word",
        },
    )

    registry = _word_registry()
    tool = create_write_document_tool(registry)
    workspace = tmp_path / "workspace"
    outputs = workspace / "outputs"
    workspace.mkdir()
    image = tmp_path / "chart.png"
    image.write_bytes(base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwC"
        "AAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
    ))
    spec = workspace / "long-report.json"
    spec.write_text(json.dumps({
        "heading_numbering": True,
        "sections": [
            {
                "page": {"size": "A4", "orientation": "portrait"},
                "footer": {"text": ""},
                "blocks": [
                    {
                        "type": "heading",
                        "level": 1,
                        "text": "年度经营报告",
                        "numbered": False,
                    },
                    {"type": "paragraph", "text": "星海科技"},
                ],
            },
            {
                "page": {"size": "A4", "orientation": "portrait"},
                "footer": {"text": "第 ", "page_number": True},
                "page_number": {"start": 1, "format": "lower_roman"},
                "blocks": [
                    {"type": "table_of_contents", "title": "目录", "levels": [1, 3]},
                ],
            },
            {
                "page": {"size": "A4", "orientation": "portrait"},
                "header": {"text": "年度经营报告"},
                "footer": {"text": "第 ", "page_number": True},
                "page_number": {"start": 1, "format": "decimal"},
                "blocks": [
                    {"type": "heading", "level": 1, "text": "经营概况"},
                    {"type": "heading", "level": 2, "text": "关键指标"},
                    {
                        "type": "image",
                        "attachment_id": "chart-1",
                        "width_cm": 4,
                        "caption": "销售趋势",
                    },
                ],
            },
            {
                "page": {
                    "size": "A4",
                    "orientation": "landscape",
                    "margins_cm": {"left": 1.5, "right": 1.5},
                },
                "header": {"text": "年度经营报告"},
                "footer": {"text": "第 ", "page_number": True},
                "blocks": [{
                    "type": "table",
                    "caption": "区域经营明细",
                    "headers": ["区域", "收入", "成本", "利润"],
                    "rows": [["华东", "100", "60", "40"]],
                }],
            },
            {
                "page": {"size": "A4", "orientation": "portrait"},
                "header": {"text": "年度经营报告"},
                "footer": {"text": "第 ", "page_number": True},
                "blocks": [
                    {"type": "heading", "level": 1, "text": "结论"},
                    {"type": "paragraph", "text": "经营保持稳定。"},
                ],
            },
        ],
    }, ensure_ascii=False), encoding="utf-8")
    attachment = {
        "id": "chart-1",
        "name": "chart.png",
        "kind": "image",
        "local_path": str(image),
    }

    with bind_tool_execution(
        tool_call_id="call-long-report",
        tool_name="write_document",
        arguments={},
        artifact_callback=None,
        attachments=(attachment,),
        workspace_root=str(workspace),
        output_root=str(outputs),
    ):
        result = tool.execute(
            format="word",
            specification_path="long-report.json",
            output_name="long-report.docx",
        )

    assert result["success"] is True
    assert result["validation"]["sections"] == 5
    assert result["validation"]["field_refresh"] == {
        "status": "updated",
        "performed": True,
        "backend": "test-word",
    }
    document = Document(outputs / "long-report.docx")
    assert [section.orientation for section in document.sections] == [
        WD_ORIENT.PORTRAIT,
        WD_ORIENT.PORTRAIT,
        WD_ORIENT.PORTRAIT,
        WD_ORIENT.LANDSCAPE,
        WD_ORIENT.PORTRAIT,
    ]
    document_xml = document._element.xml
    assert "TOC" in document_xml
    assert "SEQ XiaomeiFigure" in document_xml
    assert "SEQ XiaomeiTable" in document_xml
    assert "目录将在打开文档时更新" in document_xml
    assert "图 1  销售趋势" in [p.text for p in document.paragraphs]
    assert "表 1  区域经营明细" in [p.text for p in document.paragraphs]
    numbered_headings = [
        paragraph for paragraph in document.paragraphs
        if paragraph.style.name.startswith("Heading")
        and paragraph._p.pPr is not None
        and paragraph._p.pPr.find(qn("w:numPr")) is not None
    ]
    assert len(numbered_headings) == 3
    cover_heading = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text == "年度经营报告"
    )
    assert cover_heading._p.pPr.find(qn("w:numPr")) is None
    toc_page_number = document.sections[1]._sectPr.find(qn("w:pgNumType"))
    assert toc_page_number is not None
    assert toc_page_number.get(qn("w:start")) == "1"
    assert toc_page_number.get(qn("w:fmt")) == "lowerRoman"
    body_page_number = document.sections[2]._sectPr.find(qn("w:pgNumType"))
    assert body_page_number is not None
    assert body_page_number.get(qn("w:start")) == "1"
    assert body_page_number.get(qn("w:fmt")) == "decimal"
    assert document.sections[3]._sectPr.find(qn("w:pgNumType")) is None
    assert document.sections[4]._sectPr.find(qn("w:pgNumType")) is None


def test_write_document_rejects_workspace_image_traversal(tmp_path):
    registry = _word_registry()
    tool = create_write_document_tool(registry)
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    spec = workspace / "unsafe.json"
    spec.write_text(json.dumps({
        "blocks": [{"type": "image", "workspace_path": "../outside.png"}],
    }), encoding="utf-8")

    with bind_tool_execution(
        tool_call_id="call-unsafe-image",
        tool_name="write_document",
        arguments={},
        artifact_callback=None,
        workspace_root=str(workspace),
        output_root=str(workspace / "outputs"),
    ):
        result = tool.execute(
            format="word",
            specification_path="unsafe.json",
            output_name="unsafe.docx",
        )

    assert "relative workspace path" in result["error"]
    assert not (workspace / "outputs" / "unsafe.docx").exists()


def test_write_document_rejects_paths_and_unowned_sources(tmp_path):
    registry = _word_registry()
    tool = create_write_document_tool(registry)
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    outside = tmp_path / "outside.json"
    outside.write_text("{}", encoding="utf-8")
    valid = workspace / "valid.json"
    valid.write_text(json.dumps({
        "blocks": [{"type": "paragraph", "text": "Valid"}],
    }), encoding="utf-8")

    with bind_tool_execution(
        tool_call_id="call-3",
        tool_name="write_document",
        arguments={},
        artifact_callback=None,
        workspace_root=str(workspace),
        output_root=str(workspace / "outputs"),
    ):
        outside_result = tool.execute(
            format="word",
            specification_path=str(outside),
            output_name="result.docx",
        )
        traversal_result = tool.execute(
            format="word",
            specification_path="valid.json",
            output_name="../result.docx",
        )
        source_result = tool.execute(
            format="word",
            specification_path="valid.json",
            output_name="result.docx",
            source_attachment_id="not-owned",
        )

    assert "error" in outside_result
    assert "error" in traversal_result
    assert "error" in source_result


def test_write_document_preserves_existing_output_and_removes_failed_temporary_file(tmp_path):
    registry = _word_registry()
    tool = create_write_document_tool(registry)
    workspace = tmp_path / "workspace"
    outputs = workspace / "outputs"
    outputs.mkdir(parents=True)
    existing = outputs / "report.docx"
    existing.write_bytes(b"existing deliverable")
    valid = workspace / "valid.json"
    valid.write_text(json.dumps({
        "blocks": [{"type": "paragraph", "text": "New deliverable"}],
    }), encoding="utf-8")
    invalid = workspace / "invalid.json"
    invalid.write_text(json.dumps({
        "blocks": [{"type": "unsupported"}],
    }), encoding="utf-8")

    with bind_tool_execution(
        tool_call_id="call-safe-output",
        tool_name="write_document",
        arguments={},
        artifact_callback=None,
        workspace_root=str(workspace),
        output_root=str(outputs),
    ):
        created = tool.execute(
            format="word",
            specification_path="valid.json",
            output_name="report.docx",
        )
        failed = tool.execute(
            format="word",
            specification_path="invalid.json",
            output_name="broken.docx",
        )

    assert created["success"] is True
    assert created["output_name"] == "report (1).docx"
    assert existing.read_bytes() == b"existing deliverable"
    assert (outputs / "report (1).docx").is_file()
    assert "error" in failed
    assert not (outputs / "broken.docx").exists()
    assert not list(outputs.glob(".*.tmp.docx"))


def test_written_word_is_discovered_by_existing_artifact_pipeline(tmp_path, monkeypatch):
    monkeypatch.setattr(artifact_module.Path, "home", classmethod(lambda cls: tmp_path))
    registry = _word_registry()
    tool = create_write_document_tool(registry)
    workspace = tmp_path / ".xiaomei-brain" / "xiaomei" / "workspace"
    workspace.mkdir(parents=True)
    spec = workspace / "artifact.json"
    spec.write_text(json.dumps({
        "blocks": [{"type": "paragraph", "text": "Deliverable"}],
    }), encoding="utf-8")
    with bind_tool_execution(
        tool_call_id="call-4",
        tool_name="write_document",
        arguments={},
        artifact_callback=None,
        workspace_root=str(workspace),
        output_root=str(workspace),
    ):
        result = tool.execute(
            format="word",
            specification_path="artifact.json",
            output_name="deliverable.docx",
        )

    discovered = discover_tool_artifacts(
        "xiaomei",
        "session-1",
        "turn-1",
        "write_document",
        {"output_name": "deliverable.docx"},
        json.dumps(result, ensure_ascii=False),
    )

    assert len(discovered) == 1
    assert discovered[0]["name"] == "deliverable.docx"
    assert discovered[0]["kind"] == "document"
