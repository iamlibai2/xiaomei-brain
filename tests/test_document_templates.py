from __future__ import annotations

import json
import sqlite3
from pathlib import Path

import pytest
from docx import Document

from xiaomei_brain.documents.templates import (
    DocumentTemplateService,
    DocumentTemplateStore,
)
from xiaomei_brain.plugin.context import PluginContext
from xiaomei_brain.plugin.registry import PluginRegistry
from xiaomei_brain.plugins.tools.document_word.adapter import register as register_word
from xiaomei_brain.tools.builtin.document_templates import (
    create_manage_document_template_tool,
)
from xiaomei_brain.tools.builtin.documents import create_write_document_tool
from xiaomei_brain.tools.execution_context import bind_tool_execution


def _registry() -> PluginRegistry:
    registry = PluginRegistry()
    register_word(PluginContext({}, "document_word", "test", registry))
    return registry


def _template(path: Path, *, suffix: str = "") -> None:
    document = Document()
    document.add_heading("项目报告", level=1)
    document.add_paragraph("客户：{{customer_name}}")
    document.add_paragraph("{{BLOCK:summary}}")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "负责人"
    table.cell(0, 1).text = "{{owner}}"
    table.cell(1, 0).text = "说明"
    table.cell(1, 1).text = suffix or "固定内容"
    document.sections[0].header.paragraphs[0].text = "企业项目中心"
    document.save(path)


@pytest.fixture
def preview_stub(monkeypatch):
    def render(source: Path, output: Path):
        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_bytes(b"\x89PNG\r\n\x1a\npreview")
        return {"status": "passed", "performed": True, "backend": "test"}

    monkeypatch.setattr(
        "xiaomei_brain.documents.templates.service.render_office_preview",
        render,
    )


def test_template_store_migration_preserves_existing_database(tmp_path):
    db_path = tmp_path / "brain.db"
    connection = sqlite3.connect(db_path)
    connection.execute("CREATE TABLE sentinel (value TEXT)")
    connection.execute("INSERT INTO sentinel VALUES ('kept')")
    connection.commit()
    connection.close()

    store = DocumentTemplateStore(db_path)
    try:
        assert store._get_schema_version("document_templates") == 1
        row = store._get_conn().execute("SELECT value FROM sentinel").fetchone()
        assert row["value"] == "kept"
    finally:
        store.close()


def test_template_service_registers_analyzes_scopes_and_removes(
    tmp_path,
    preview_stub,
):
    source = tmp_path / "company-template.docx"
    _template(source)
    agent_root = tmp_path / "agent"
    service = DocumentTemplateService(
        _registry(),
        agent_root,
        agent_root / "memory" / "brain.db",
    )

    personal = service.register(
        source,
        name="项目周报",
        person_id="person-a",
        description="项目每周汇报",
        keywords=["周报", "项目进展", "周报"],
    )

    assert personal.scope_type == "person"
    assert personal.keywords == ("周报", "项目进展")
    assert {item["key"] for item in personal.manifest["placeholders"]} == {
        "customer_name",
        "owner",
        "summary",
    }
    assert next(
        item for item in personal.manifest["placeholders"]
        if item["key"] == "summary"
    )["type"] == "blocks"
    assert service.list("person-b") == []

    global_template = service.register(
        source,
        name="公司项目报告",
        person_id="person-a",
        scope_type="global",
    )
    assert [item.template_id for item in service.list("person-b")] == [
        global_template.template_id,
    ]
    _, stored_source = service.source_for_use(personal.template_id, "person-a")
    assert stored_source.read_bytes() == source.read_bytes()
    with pytest.raises(ValueError, match="没有找到"):
        service.source_for_use(personal.template_id, "person-b")

    output_root = agent_root / "workspace" / "work"
    preview = service.copy_preview_to(personal, output_root, "session-a")
    assert preview is not None and preview.is_file()
    replacement = tmp_path / "replacement.docx"
    _template(replacement, suffix="新版模板")
    updated = service.update(
        personal.template_id,
        "person-a",
        source_path=replacement,
        name="研发项目周报",
        keywords=["研发", "周报"],
    )
    assert updated.name == "研发项目周报"
    assert updated.sha256 != personal.sha256
    assert updated.keywords == ("研发", "周报")
    removed = service.remove(personal.template_id, "person-a")
    assert removed.template_id == personal.template_id
    assert not stored_source.exists()


def test_conversational_template_tool_and_writer_use_the_same_service(
    tmp_path,
    preview_stub,
):
    registry = _registry()
    agent_root = tmp_path / "agent"
    workspace = agent_root / "workspace"
    work = workspace / "work"
    outputs = workspace / "outputs"
    work.mkdir(parents=True)
    outputs.mkdir(parents=True)
    source = tmp_path / "weekly.docx"
    _template(source)
    service = DocumentTemplateService(
        registry,
        agent_root,
        agent_root / "memory" / "brain.db",
    )
    manage_tool = create_manage_document_template_tool(service)

    context_args = {
        "tool_call_id": "call-register",
        "tool_name": manage_tool.name,
        "arguments": {},
        "artifact_callback": None,
        "session_id": "session-a",
        "person_id": "person-a",
        "attachments": ({
            "id": "attachment-1",
            "kind": "document",
            "name": source.name,
            "local_path": str(source),
        },),
        "workspace_root": str(workspace),
        "working_directory": str(work),
        "output_root": str(outputs),
    }
    with bind_tool_execution(**context_args):
        registered = manage_tool.execute(
            action="register",
            attachment_id="attachment-1",
            name="项目周报",
            description="项目进展模板",
            keywords=["周报"],
            scope_type="person",
        )
    template_id = registered["template"]["template_id"]
    assert Path(registered["preview_output_path"]).is_file()

    specification_path = work / "report.json"
    specification_path.write_text(json.dumps({
        "operations": [
            {
                "type": "replace_placeholders",
                "values": {"customer_name": "星海科技", "owner": "李白"},
            },
            {
                "type": "insert_blocks_after",
                "marker": "{{BLOCK:summary}}",
                "remove_marker": True,
                "blocks": [
                    {"type": "heading", "level": 2, "text": "本周进展"},
                    {"type": "paragraph", "text": "核心功能已经完成。"},
                ],
            },
        ],
    }, ensure_ascii=False), encoding="utf-8")
    write_tool = create_write_document_tool(
        registry,
        template_service=service,
    )
    with bind_tool_execution(
        **{
            **context_args,
            "tool_call_id": "call-write",
            "tool_name": write_tool.name,
            "attachments": (),
        },
    ):
        result = write_tool.execute(
            format="word",
            specification_path="work/report.json",
            output_name="项目周报.docx",
            template_id=template_id,
        )

    assert result["success"] is True
    assert result["template"] == {
        "template_id": template_id,
        "name": "项目周报",
    }
    generated = Document(result["output_path"])
    generated_text = "\n".join(
        [paragraph.text for paragraph in generated.paragraphs]
        + [cell.text for table in generated.tables for row in table.rows for cell in row.cells]
    )
    assert "星海科技" in generated_text
    assert "核心功能已经完成" in generated_text
    assert "{{" not in generated_text
    original = Document(source)
    assert "{{customer_name}}" in "\n".join(p.text for p in original.paragraphs)

    incomplete_path = work / "incomplete.json"
    incomplete_path.write_text(json.dumps({
        "operations": [{
            "type": "replace_placeholders",
            "values": {"customer_name": "星海科技", "owner": "李白"},
        }],
    }, ensure_ascii=False), encoding="utf-8")
    with bind_tool_execution(
        **{
            **context_args,
            "tool_call_id": "call-incomplete",
            "tool_name": write_tool.name,
            "attachments": (),
        },
    ):
        incomplete = write_tool.execute(
            format="word",
            specification_path="work/incomplete.json",
            output_name="不完整周报.docx",
            template_id=template_id,
        )
    assert "未填写字段" in incomplete["error"]
    assert not (outputs / "不完整周报.docx").exists()


def test_remove_tool_requires_explicit_confirmation(tmp_path, preview_stub):
    registry = _registry()
    agent_root = tmp_path / "agent"
    workspace = agent_root / "workspace"
    workspace.mkdir(parents=True)
    source = tmp_path / "template.docx"
    _template(source)
    service = DocumentTemplateService(
        registry,
        agent_root,
        agent_root / "memory" / "brain.db",
    )
    record = service.register(source, name="合同", person_id="person-a")
    tool = create_manage_document_template_tool(service)
    with bind_tool_execution(
        tool_call_id="call-remove",
        tool_name=tool.name,
        arguments={},
        artifact_callback=None,
        session_id="session-a",
        person_id="person-a",
        workspace_root=str(workspace),
        working_directory=str(workspace),
        output_root=str(workspace),
    ):
        denied = tool.execute(action="remove", template=record.template_id)
        removed = tool.execute(
            action="remove",
            template=record.template_id,
            confirmed=True,
        )
    assert "confirmed=true" in denied["error"]
    assert removed["success"] is True
