"""Low-level Word structures that python-docx does not expose directly."""

from __future__ import annotations

import base64
import os
import shutil
import subprocess
import sys
from pathlib import Path
from typing import Any


_WORD_REFRESH_SCRIPT = r"""
$source = [System.IO.Path]::GetFullPath($env:XIAOMEI_WORD_REFRESH_SOURCE)
$application = $null
$document = $null
try {
    $application = New-Object -ComObject Word.Application
    $application.Visible = $false
    $application.DisplayAlerts = 0
    try { $application.AutomationSecurity = 3 } catch {}
    try { $application.Options.UpdateLinksAtOpen = $false } catch {}
    $document = $application.Documents.Open($source, $false, $false)
    $document.Repaginate()
    foreach ($toc in $document.TablesOfContents) {
        $toc.Update()
    }
    [void]$document.Fields.Update()
    foreach ($section in $document.Sections) {
        foreach ($header in $section.Headers) {
            [void]$header.Range.Fields.Update()
        }
        foreach ($footer in $section.Footers) {
            [void]$footer.Range.Fields.Update()
        }
    }
    $document.Repaginate()
    $document.Save()
}
finally {
    if ($null -ne $document) {
        try { $document.Close($false) } catch {}
        [void][System.Runtime.InteropServices.Marshal]::FinalReleaseComObject($document)
    }
    if ($null -ne $application) {
        try { $application.Quit() } catch {}
        [void][System.Runtime.InteropServices.Marshal]::FinalReleaseComObject($application)
    }
    [GC]::Collect()
    [GC]::WaitForPendingFinalizers()
}
"""


def refresh_word_fields(
    source_path: str | Path,
    *,
    timeout: float = 60.0,
) -> dict[str, Any]:
    """Refresh paginated Word fields when desktop Microsoft Word is available."""
    source = Path(source_path).resolve(strict=True)
    if source.suffix.lower() != ".docx":
        raise ValueError("只能刷新 DOCX 文档域")
    if sys.platform != "win32":
        return {
            "status": "unavailable",
            "performed": False,
            "reason": "当前平台不支持 Microsoft Word COM",
        }
    executable = shutil.which("powershell.exe") or shutil.which("powershell")
    if not executable:
        return {
            "status": "unavailable",
            "performed": False,
            "reason": "PowerShell 不可用",
        }
    encoded = base64.b64encode(
        _WORD_REFRESH_SCRIPT.encode("utf-16-le")
    ).decode("ascii")
    environment = os.environ.copy()
    environment["XIAOMEI_WORD_REFRESH_SOURCE"] = str(source)
    creation_flags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
    try:
        completed = subprocess.run(
            [
                executable,
                "-NoLogo",
                "-NoProfile",
                "-NonInteractive",
                "-EncodedCommand",
                encoded,
            ],
            capture_output=True,
            text=True,
            errors="replace",
            env=environment,
            timeout=timeout,
            check=False,
            creationflags=creation_flags,
        )
    except (OSError, subprocess.SubprocessError) as exc:
        return {
            "status": "unavailable",
            "performed": False,
            "reason": str(exc)[-500:],
        }
    if completed.returncode != 0:
        detail = (completed.stderr or completed.stdout or "Word 域更新失败").strip()
        return {
            "status": "unavailable",
            "performed": False,
            "reason": detail[-500:],
        }
    return {
        "status": "updated",
        "performed": True,
        "backend": "microsoft-word-com",
    }


def enable_field_updates(document: Any) -> None:
    """Ask Word-compatible applications to refresh TOC and sequence fields."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_field(
    paragraph: Any,
    instruction: str,
    *,
    cached_text: str = "",
) -> None:
    """Append a complex Word field with a useful pre-refresh fallback value."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = f" {instruction.strip()} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction_node, separate])
    if cached_text:
        text = OxmlElement("w:t")
        text.text = cached_text
        run._r.append(text)
    run._r.append(end)


def add_table_of_contents(
    document: Any,
    *,
    levels: tuple[int, int] = (1, 3),
    title: str = "目录",
) -> list[Any]:
    """Insert a native TOC field based on Word heading styles."""
    start, end = levels
    if not 1 <= start <= end <= 9:
        raise ValueError("table_of_contents.levels 必须在 1 到 9 之间")
    paragraphs: list[Any] = []
    if title:
        # Title style is intentionally outside the TOC's Heading 1-9 range.
        paragraphs.append(document.add_heading(title, level=0))
    paragraph = document.add_paragraph()
    add_field(
        paragraph,
        f'TOC \\o "{start}-{end}" \\h \\z \\u',
        cached_text="目录将在打开文档时更新",
    )
    enable_field_updates(document)
    paragraphs.append(paragraph)
    return paragraphs


def ensure_heading_numbering(document: Any) -> int:
    """Create one reusable 1 / 1.1 / 1.1.1 multilevel numbering definition."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "multilevel")
    abstract.append(multi_level)

    for level in range(9):
        level_node = OxmlElement("w:lvl")
        level_node.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        number_format = OxmlElement("w:numFmt")
        number_format.set(qn("w:val"), "decimal")
        level_text = OxmlElement("w:lvlText")
        level_text.set(
            qn("w:val"),
            ".".join(f"%{index}" for index in range(1, level + 2)),
        )
        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "space")
        paragraph_properties = OxmlElement("w:pPr")
        indentation = OxmlElement("w:ind")
        indentation.set(qn("w:left"), str(360 * level))
        indentation.set(qn("w:hanging"), "0")
        paragraph_properties.append(indentation)
        level_node.extend(
            [start, number_format, level_text, suffix, paragraph_properties]
        )
        abstract.append(level_node)

    numbering.append(abstract)
    concrete = OxmlElement("w:num")
    concrete.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    concrete.append(abstract_reference)
    numbering.append(concrete)
    return num_id


def apply_heading_number(paragraph: Any, *, level: int, num_id: int) -> None:
    """Attach a heading paragraph to the document's multilevel numbering."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = paragraph._p.get_or_add_pPr()
    existing = properties.find(qn("w:numPr"))
    if existing is not None:
        properties.remove(existing)
    numbering_properties = OxmlElement("w:numPr")
    indentation_level = OxmlElement("w:ilvl")
    indentation_level.set(qn("w:val"), str(max(0, min(level - 1, 8))))
    number_reference = OxmlElement("w:numId")
    number_reference.set(qn("w:val"), str(num_id))
    numbering_properties.extend([indentation_level, number_reference])
    properties.append(numbering_properties)


def add_sequence_caption(
    document: Any,
    *,
    kind: str,
    text: str,
    number: int,
    alignment: Any | None = None,
) -> Any:
    """Add a native, refreshable figure or table caption."""
    labels = {
        "figure": ("图", "XiaomeiFigure"),
        "table": ("表", "XiaomeiTable"),
    }
    if kind not in labels:
        raise ValueError("caption.kind 仅支持 figure 或 table")
    label, sequence = labels[kind]
    paragraph = document.add_paragraph(style="Caption")
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.add_run(f"{label} ")
    add_field(paragraph, f"SEQ {sequence} \\* ARABIC", cached_text=str(number))
    if text:
        paragraph.add_run(f"  {text}")
    enable_field_updates(document)
    return paragraph


def configure_page_numbering(
    section: Any,
    *,
    start: int | None = None,
    number_format: str | None = None,
) -> None:
    """Configure page-number restart and display format for one section."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    formats = {
        "decimal": "decimal",
        "lower_roman": "lowerRoman",
        "upper_roman": "upperRoman",
    }
    if number_format is not None and number_format not in formats:
        raise ValueError(
            "page_number.format 仅支持 decimal、lower_roman 或 upper_roman"
        )
    if start is not None and not 1 <= int(start) <= 9999:
        raise ValueError("page_number.start 必须在 1 到 9999 之间")
    section_properties = section._sectPr
    page_number = section_properties.find(qn("w:pgNumType"))
    if page_number is None:
        page_number = OxmlElement("w:pgNumType")
        section_properties.append(page_number)
    if start is not None:
        page_number.set(qn("w:start"), str(int(start)))
    if number_format is not None:
        page_number.set(qn("w:fmt"), formats[number_format])
