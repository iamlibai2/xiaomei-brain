"""Deterministic DOCX creation and common non-destructive revisions."""

from __future__ import annotations

import shutil
from pathlib import Path
from typing import Any, Iterable

from xiaomei_brain.documents import render_office_document
from xiaomei_brain.plugins.tools.document_word.extractor import WordExtractor
from xiaomei_brain.plugins.tools.document_word.template_analyzer import (
    WordTemplateAnalyzer,
)
from xiaomei_brain.plugins.tools.document_word.structure import (
    add_sequence_caption,
    add_table_of_contents,
    apply_heading_number,
    configure_page_numbering,
    ensure_heading_numbering,
    refresh_word_fields,
)
from xiaomei_brain.plugins.tools.document_word.theme import (
    apply_word_theme,
    resolve_word_theme,
    style_header_footer,
    style_table_cells,
    style_word_table,
)


class WordWriter:
    format_id = "word"
    suffix = ".docx"
    writer_version = "1.4.0"
    # Template inspection is an optional capability of this format writer,
    # not a separate plugin category.
    template_analyzer = WordTemplateAnalyzer()

    def write(
        self,
        specification: dict[str, Any],
        output_path: Path,
        *,
        source_path: Path | None = None,
        asset_paths: dict[str, Path] | None = None,
    ) -> dict[str, Any]:
        try:
            from docx import Document
        except ImportError as exc:
            raise ValueError("Word 写入依赖 python-docx 未安装") from exc

        resolved_theme = resolve_word_theme(specification.get("theme"))
        structure_state: dict[str, Any] = {
            "caption_counts": {"figure": 0, "table": 0},
            "heading_num_id": None,
            "requires_pagination": False,
        }
        heading_numbering = specification.get("heading_numbering") is True
        if source_path is not None:
            if source_path.suffix.lower() != self.suffix:
                raise ValueError("Word writer 只能修改 DOCX 附件")
            shutil.copy2(source_path, output_path)
            document = Document(str(output_path))
            # Existing documents and enterprise templates own their visual
            # language. Only an explicit theme request may restyle them.
            theme = resolved_theme if "theme" in specification else None
            if theme is not None:
                apply_word_theme(document, theme)
            self._set_page_layout(document, specification.get("page"))
            operations = specification.get("operations")
            if not isinstance(operations, list) or not operations:
                raise ValueError("修改 Word 文档时 specification.operations 不能为空")
            replacements = 0
            insertions = 0
            styled_cells = 0
            for operation in operations:
                if not isinstance(operation, dict):
                    raise ValueError("Word operation 必须是对象")
                kind = str(operation.get("type") or "")
                if kind == "replace_text":
                    old = str(operation.get("old") or "")
                    new = str(operation.get("new") or "")
                    if not old:
                        raise ValueError("replace_text.old 不能为空")
                    replacements += self._replace_text(
                        document,
                        old,
                        new,
                        replace_all=operation.get("all") is True,
                    )
                elif kind == "replace_placeholders":
                    values = operation.get("values")
                    if not isinstance(values, dict) or not values:
                        raise ValueError("replace_placeholders.values 必须是非空对象")
                    missing: list[str] = []
                    for key, value in values.items():
                        placeholder = str(key)
                        if not (
                            placeholder.startswith("{{")
                            and placeholder.endswith("}}")
                        ):
                            placeholder = "{{" + placeholder + "}}"
                        count = self._replace_text(
                            document,
                            placeholder,
                            str(value),
                            replace_all=True,
                            required=False,
                        )
                        replacements += count
                        if count == 0:
                            missing.append(placeholder)
                    if missing and operation.get("allow_missing") is not True:
                        raise ValueError(
                            "Word 文档中没有找到占位符: " + ", ".join(missing)
                        )
                elif kind == "append_blocks":
                    blocks = operation.get("blocks")
                    if not isinstance(blocks, list):
                        raise ValueError("append_blocks.blocks 必须是数组")
                    insertions += len(self._append_blocks(
                        document,
                        blocks,
                        asset_paths=asset_paths,
                        theme=theme,
                        structure_state=structure_state,
                        heading_numbering=heading_numbering,
                    ))
                elif kind == "insert_blocks_after":
                    marker = str(operation.get("marker") or "")
                    blocks = operation.get("blocks")
                    if not marker:
                        raise ValueError("insert_blocks_after.marker 不能为空")
                    if not isinstance(blocks, list) or not blocks:
                        raise ValueError("insert_blocks_after.blocks 必须是非空数组")
                    insertions += self._insert_blocks_after(
                        document,
                        marker,
                        blocks,
                        remove_marker=operation.get("remove_marker") is True,
                        asset_paths=asset_paths,
                        theme=theme,
                        structure_state=structure_state,
                        heading_numbering=heading_numbering,
                    )
                elif kind == "set_page_layout":
                    self._set_page_layout(document, operation)
                elif kind == "set_header_footer":
                    self._set_header_footer(
                        document,
                        operation.get("header"),
                        operation.get("footer"),
                        theme=theme,
                    )
                elif kind == "set_properties":
                    self._set_properties(document, operation)
                elif kind == "style_table_cells":
                    table_index = operation.get("table_index")
                    if (
                        isinstance(table_index, bool)
                        or not isinstance(table_index, int)
                        or table_index < 1
                        or table_index > len(document.tables)
                    ):
                        raise ValueError(
                            "style_table_cells.table_index 超出范围；"
                            f"当前文档共有 {len(document.tables)} 张表格"
                        )
                    styled_cells += style_table_cells(
                        document.tables[table_index - 1],
                        operation,
                    )
                else:
                    raise ValueError(f"不支持的 Word operation: {kind}")
            document.save(str(output_path))
        else:
            theme = resolved_theme
            document = Document()
            self._set_properties(document, specification.get("properties", {}))
            apply_word_theme(document, theme)
            self._set_default_style(document, specification.get("default_style"))
            self._set_page_layout(
                document,
                specification.get("page") or {"size": "A4"},
            )
            self._set_header_footer(
                document,
                specification.get("header"),
                specification.get("footer"),
                theme=theme,
            )
            title = str(specification.get("title") or "").strip()
            if title:
                document.add_heading(title, level=0)
            subtitle = str(specification.get("subtitle") or "").strip()
            if subtitle:
                document.add_paragraph(subtitle, style="Subtitle")
            sections = specification.get("sections")
            blocks = specification.get("blocks")
            if sections is not None:
                if blocks is not None:
                    raise ValueError("创建 Word 文档时 blocks 和 sections 不能同时使用")
                if not isinstance(sections, list) or not sections:
                    raise ValueError("specification.sections 必须是非空数组")
                self._append_sections(
                    document,
                    sections,
                    asset_paths=asset_paths,
                    theme=theme,
                    structure_state=structure_state,
                    heading_numbering=heading_numbering,
                )
            else:
                if not isinstance(blocks, list) or not blocks:
                    raise ValueError("创建 Word 文档时 specification.blocks 不能为空")
                self._append_blocks(
                    document,
                    blocks,
                    asset_paths=asset_paths,
                    theme=theme,
                    structure_state=structure_state,
                    heading_numbering=heading_numbering,
                )
            document.save(str(output_path))
            replacements = 0
            insertions = 0
            styled_cells = 0

        field_refresh = {
            "status": "not_required",
            "performed": False,
        }
        if (
            structure_state.get("requires_pagination") is True
            and specification.get("refresh_fields") is not False
        ):
            field_refresh = refresh_word_fields(output_path)

        # Structural and semantic verification happen before the file is
        # announced as an artifact.
        verified = Document(str(output_path))
        if not verified.paragraphs and not verified.tables:
            raise ValueError("生成的 Word 文档为空")
        extraction = WordExtractor().extract(output_path)
        preview = extraction.sections[0].content[:1200] if extraction.sections else ""
        render_validation = (
            render_office_document(output_path)
            if specification.get("visual_validation") is True
            else {
                "status": "disabled",
                "performed": False,
                "reason": "specification.visual_validation 未启用",
            }
        )
        return {
            "writer": self.format_id,
            "writer_version": self.writer_version,
            "validation": {
                "valid": True,
                "paragraphs": len(verified.paragraphs),
                "tables": len(verified.tables),
                "sections": len(verified.sections),
                "replacements": replacements,
                "insertions": insertions,
                "styled_cells": styled_cells,
                "content_preview": preview,
                "theme": theme["preset"] if theme is not None else "preserved",
                "field_refresh": field_refresh,
                "render_validation": render_validation,
            },
        }

    @staticmethod
    def _set_properties(document: Any, values: Any) -> None:
        if not isinstance(values, dict):
            return
        properties = document.core_properties
        for key in ("title", "subject", "author", "keywords", "comments"):
            if key in values:
                setattr(properties, key, str(values[key]))

    @staticmethod
    def _set_default_style(document: Any, values: Any) -> None:
        if not isinstance(values, dict):
            return
        from docx.oxml.ns import qn
        from docx.shared import Pt

        style = document.styles["Normal"]
        font_name = str(values.get("font") or "").strip()
        if font_name:
            style.font.name = font_name
            fonts = style._element.get_or_add_rPr().rFonts
            for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
                fonts.set(qn(f"w:{attribute}"), font_name)
            for attribute in (
                "asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"
            ):
                qualified = qn(f"w:{attribute}")
                if qualified in fonts.attrib:
                    del fonts.attrib[qualified]
        if values.get("size_pt") is not None:
            size = float(values["size_pt"])
            if not 6 <= size <= 72:
                raise ValueError("default_style.size_pt 必须在 6 到 72 之间")
            style.font.size = Pt(size)

    @staticmethod
    def _set_page_layout(document: Any, values: Any) -> None:
        if not isinstance(values, dict) or not values:
            return
        for section in document.sections:
            WordWriter._set_section_layout(section, values)

    @staticmethod
    def _set_section_layout(section: Any, values: Any) -> None:
        """Apply page geometry to one section without touching its siblings."""
        if not isinstance(values, dict) or not values:
            return
        from docx.enum.section import WD_ORIENT
        from docx.shared import Cm

        size_name = str(values.get("size") or "").strip().upper()
        sizes = {
            "A4": (21.0, 29.7),
            "LETTER": (21.59, 27.94),
        }
        if size_name and size_name not in sizes:
            raise ValueError("page.size 仅支持 A4 或 Letter")
        orientation = str(values.get("orientation") or "portrait").lower()
        if orientation not in {"portrait", "landscape"}:
            raise ValueError("page.orientation 仅支持 portrait 或 landscape")
        margins = values.get("margins_cm", {})
        if margins is not None and not isinstance(margins, dict):
            raise ValueError("page.margins_cm 必须是对象")

        if size_name:
            width, height = sizes[size_name]
            if orientation == "landscape":
                width, height = height, width
            section.page_width = Cm(width)
            section.page_height = Cm(height)
        elif (
            orientation == "landscape"
            and section.page_width < section.page_height
        ) or (
            orientation == "portrait"
            and section.page_width > section.page_height
        ):
            section.page_width, section.page_height = (
                section.page_height,
                section.page_width,
            )
        section.orientation = (
            WD_ORIENT.LANDSCAPE
            if orientation == "landscape"
            else WD_ORIENT.PORTRAIT
        )
        for key, attribute in (
            ("top", "top_margin"),
            ("right", "right_margin"),
            ("bottom", "bottom_margin"),
            ("left", "left_margin"),
        ):
            if key not in margins:
                continue
            margin = float(margins[key])
            if not 0 <= margin <= 10:
                raise ValueError(f"page.margins_cm.{key} 必须在 0 到 10 之间")
            setattr(section, attribute, Cm(margin))

    @staticmethod
    def _add_page_number(paragraph: Any) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = " PAGE "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend([begin, instruction, end])

    @classmethod
    def _set_header_footer(
        cls,
        document: Any,
        header_values: Any,
        footer_values: Any,
        *,
        theme: dict[str, Any] | None,
    ) -> None:
        if header_values is not None and not isinstance(header_values, dict):
            raise ValueError("header 必须是对象")
        if footer_values is not None and not isinstance(footer_values, dict):
            raise ValueError("footer 必须是对象")
        for section in document.sections:
            cls._set_section_header_footer(
                section,
                header_values,
                footer_values,
                theme=theme,
                unlink=False,
            )

    @classmethod
    def _set_section_header_footer(
        cls,
        section: Any,
        header_values: Any,
        footer_values: Any,
        *,
        theme: dict[str, Any] | None,
        unlink: bool,
    ) -> None:
        """Set one section's running content, optionally breaking inheritance."""
        if header_values is not None and not isinstance(header_values, dict):
            raise ValueError("section.header 必须是对象")
        if footer_values is not None and not isinstance(footer_values, dict):
            raise ValueError("section.footer 必须是对象")
        if isinstance(header_values, dict):
            if unlink:
                section.header.is_linked_to_previous = False
            paragraph = section.header.paragraphs[0]
            paragraph.clear()
            paragraph.add_run(str(header_values.get("text") or ""))
            if theme is not None:
                style_header_footer(paragraph, theme, footer=False)
        if isinstance(footer_values, dict):
            if unlink:
                section.footer.is_linked_to_previous = False
            paragraph = section.footer.paragraphs[0]
            paragraph.clear()
            prefix = str(footer_values.get("text") or "")
            if prefix:
                paragraph.add_run(prefix)
            if footer_values.get("page_number") is True:
                cls._add_page_number(paragraph)
            if theme is not None:
                style_header_footer(paragraph, theme, footer=True)

    def _append_sections(
        self,
        document: Any,
        sections: Iterable[Any],
        *,
        asset_paths: dict[str, Path] | None,
        theme: dict[str, Any] | None,
        structure_state: dict[str, Any],
        heading_numbering: bool,
    ) -> None:
        """Build independently formatted cover, TOC, body or wide-table sections."""
        from docx.enum.section import WD_SECTION

        start_types = {
            "new_page": WD_SECTION.NEW_PAGE,
            "continuous": WD_SECTION.CONTINUOUS,
            "even_page": WD_SECTION.EVEN_PAGE,
            "odd_page": WD_SECTION.ODD_PAGE,
        }
        for index, values in enumerate(sections):
            if not isinstance(values, dict):
                raise ValueError("sections 中每一项必须是对象")
            if index == 0:
                section = document.sections[0]
            else:
                start = str(values.get("start") or "new_page").strip().lower()
                if start not in start_types:
                    raise ValueError(
                        "section.start 仅支持 new_page、continuous、even_page 或 odd_page"
                    )
                section = document.add_section(start_types[start])

            self._set_section_layout(
                section,
                values.get("page") or {"size": "A4"},
            )
            if values.get("different_first_page") is True:
                section.different_first_page_header_footer = True
            self._set_section_header_footer(
                section,
                values.get("header"),
                values.get("footer"),
                theme=theme,
                unlink=index > 0,
            )
            page_number = values.get("page_number")
            if page_number is not None:
                if not isinstance(page_number, dict):
                    raise ValueError("section.page_number 必须是对象")
                configure_page_numbering(
                    section,
                    start=page_number.get("start"),
                    number_format=page_number.get("format"),
                )
            elif index > 0:
                # python-docx clones the previous section properties. Remove
                # a cloned restart instruction so numbering naturally
                # continues across ordinary body and landscape sections.
                from docx.oxml.ns import qn

                inherited = section._sectPr.find(qn("w:pgNumType"))
                if inherited is not None:
                    section._sectPr.remove(inherited)
            blocks = values.get("blocks")
            if not isinstance(blocks, list) or not blocks:
                raise ValueError("section.blocks 必须是非空数组")
            self._append_blocks(
                document,
                blocks,
                asset_paths=asset_paths,
                theme=theme,
                structure_state=structure_state,
                heading_numbering=heading_numbering,
            )

    def _append_blocks(
        self,
        document: Any,
        blocks: Iterable[Any],
        *,
        asset_paths: dict[str, Path] | None = None,
        theme: dict[str, Any] | None,
        structure_state: dict[str, Any] | None = None,
        heading_numbering: bool = False,
    ) -> list[Any]:
        structure_state = structure_state or {
            "caption_counts": {"figure": 0, "table": 0},
            "heading_num_id": None,
            "requires_pagination": False,
        }
        elements: list[Any] = []
        for block in blocks:
            if not isinstance(block, dict):
                raise ValueError("Word block 必须是对象")
            kind = str(block.get("type") or "paragraph")
            if kind == "table_of_contents":
                levels = block.get("levels", [1, 3])
                if not isinstance(levels, list) or len(levels) != 2:
                    raise ValueError("table_of_contents.levels 必须是两个整数")
                toc_paragraphs = add_table_of_contents(
                    document,
                    levels=(int(levels[0]), int(levels[1])),
                    title=str(block.get("title") or "目录"),
                )
                structure_state["requires_pagination"] = True
                elements.extend(paragraph._p for paragraph in toc_paragraphs)
            elif kind == "heading":
                level = max(1, min(int(block.get("level", 1)), 9))
                paragraph = document.add_heading(
                    str(block.get("text") or ""),
                    level=level,
                )
                numbered = block.get("numbered")
                if numbered is True or (numbered is None and heading_numbering):
                    num_id = structure_state.get("heading_num_id")
                    if num_id is None:
                        num_id = ensure_heading_numbering(document)
                        structure_state["heading_num_id"] = num_id
                    apply_heading_number(paragraph, level=level, num_id=num_id)
                elements.append(paragraph._p)
            elif kind == "paragraph":
                paragraph = document.add_paragraph(str(block.get("text") or ""))
                style = str(block.get("style") or "").strip()
                if style:
                    try:
                        paragraph.style = style
                    except KeyError as exc:
                        raise ValueError(f"Word 样式不存在: {style}") from exc
                elements.append(paragraph._p)
            elif kind == "list":
                items = block.get("items")
                if not isinstance(items, list):
                    raise ValueError("list.items 必须是数组")
                style = "List Number" if block.get("ordered") is True else "List Bullet"
                for item in items:
                    paragraph = document.add_paragraph(str(item), style=style)
                    elements.append(paragraph._p)
            elif kind == "table":
                headers = block.get("headers", [])
                rows = block.get("rows", [])
                if not isinstance(headers, list) or not isinstance(rows, list):
                    raise ValueError("table.headers 和 table.rows 必须是数组")
                width = len(headers) or max(
                    (len(row) for row in rows if isinstance(row, list)),
                    default=0,
                )
                if width <= 0:
                    raise ValueError("表格至少需要一列")
                caption = str(block.get("caption") or "").strip()
                if caption:
                    counts = structure_state["caption_counts"]
                    counts["table"] += 1
                    caption_paragraph = add_sequence_caption(
                        document,
                        kind="table",
                        text=caption,
                        number=counts["table"],
                    )
                    elements.append(caption_paragraph._p)
                table = document.add_table(rows=1 if headers else 0, cols=width)
                explicit_style = str(block.get("style") or "").strip()
                if explicit_style:
                    table.style = explicit_style
                if headers:
                    for index, value in enumerate(headers[:width]):
                        table.rows[0].cells[index].text = str(value)
                for row in rows:
                    if not isinstance(row, list):
                        raise ValueError("table.rows 中每一项必须是数组")
                    cells = table.add_row().cells
                    for index, value in enumerate(row[:width]):
                        cells[index].text = str(value)
                if theme is not None:
                    style_word_table(
                        table,
                        theme,
                        has_header=bool(headers),
                        column_widths_cm=block.get("column_widths_cm"),
                    )
                elif block.get("column_widths_cm") is not None:
                    from docx.shared import Cm

                    widths = block["column_widths_cm"]
                    if not isinstance(widths, list) or len(widths) != width:
                        raise ValueError(
                            "table.column_widths_cm 必须与表格列数一致"
                        )
                    table.autofit = False
                    for row in table.rows:
                        for index, cell in enumerate(row.cells):
                            value = float(widths[index])
                            if not 0 < value <= 50:
                                raise ValueError(
                                    "table.column_widths_cm 每列必须大于 0 且不超过 50"
                                )
                            cell.width = Cm(value)
                elements.append(table._tbl)
            elif kind == "page_break":
                paragraph = document.add_page_break()
                elements.append(paragraph._p)
            elif kind == "quote":
                paragraph = document.add_paragraph(
                    str(block.get("text") or ""),
                    style="Quote",
                )
                elements.append(paragraph._p)
            elif kind == "image":
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.shared import Cm

                attachment_id = str(block.get("attachment_id") or "").strip()
                workspace_path = str(block.get("workspace_path") or "").strip()
                if bool(attachment_id) == bool(workspace_path):
                    raise ValueError(
                        "image 必须且只能提供 attachment_id 或 workspace_path 之一"
                    )
                asset_key = attachment_id or f"workspace:{workspace_path}"
                image_path = (asset_paths or {}).get(asset_key)
                if image_path is None or not image_path.is_file():
                    raise ValueError(
                        f"当前执行现场没有可用图片: {attachment_id or workspace_path}"
                    )
                paragraph = document.add_paragraph()
                alignment = str(block.get("align") or "center").lower()
                alignments = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                }
                if alignment not in alignments:
                    raise ValueError("image.align 仅支持 left、center 或 right")
                paragraph.alignment = alignments[alignment]
                picture_args: dict[str, Any] = {}
                if block.get("width_cm") is not None:
                    width_cm = float(block["width_cm"])
                    if not 0 < width_cm <= 50:
                        raise ValueError("image.width_cm 必须大于 0 且不超过 50")
                    picture_args["width"] = Cm(width_cm)
                if block.get("height_cm") is not None:
                    height_cm = float(block["height_cm"])
                    if not 0 < height_cm <= 50:
                        raise ValueError("image.height_cm 必须大于 0 且不超过 50")
                    picture_args["height"] = Cm(height_cm)
                paragraph.add_run().add_picture(str(image_path), **picture_args)
                elements.append(paragraph._p)
                caption = str(block.get("caption") or "").strip()
                if caption:
                    counts = structure_state["caption_counts"]
                    counts["figure"] += 1
                    caption_paragraph = add_sequence_caption(
                        document,
                        kind="figure",
                        text=caption,
                        number=counts["figure"],
                        alignment=alignments[alignment],
                    )
                    elements.append(caption_paragraph._p)
            else:
                raise ValueError(f"不支持的 Word block: {kind}")
        return elements

    @classmethod
    def _table_paragraphs(cls, tables: Iterable[Any]) -> Iterable[Any]:
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
                    yield from cls._table_paragraphs(cell.tables)

    @classmethod
    def _paragraphs(cls, document: Any) -> Iterable[Any]:
        """Yield body, table, header and footer paragraphs exactly once."""
        seen: set[Any] = set()
        containers: list[Any] = [document]
        for section in document.sections:
            containers.extend([
                section.header,
                section.first_page_header,
                section.even_page_header,
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            ])
        for container in containers:
            paragraphs = list(container.paragraphs)
            paragraphs.extend(cls._table_paragraphs(container.tables))
            for paragraph in paragraphs:
                element = paragraph._p
                if element in seen:
                    continue
                seen.add(element)
                yield paragraph

    def _insert_blocks_after(
        self,
        document: Any,
        marker: str,
        blocks: Iterable[Any],
        *,
        remove_marker: bool,
        asset_paths: dict[str, Path] | None = None,
        theme: dict[str, Any] | None,
        structure_state: dict[str, Any] | None = None,
        heading_numbering: bool = False,
    ) -> int:
        paragraph = next(
            (item for item in document.paragraphs if marker in item.text),
            None,
        )
        if paragraph is None:
            raise ValueError(f"Word 正文中没有找到插入标记: {marker}")
        elements = self._append_blocks(
            document,
            blocks,
            asset_paths=asset_paths,
            theme=theme,
            structure_state=structure_state,
            heading_numbering=heading_numbering,
        )
        anchor = paragraph._p
        for element in elements:
            anchor.addnext(element)
            anchor = element
        if remove_marker:
            self._replace_in_paragraph(
                paragraph,
                marker,
                "",
                replace_all=True,
            )
            if not paragraph.text.strip():
                parent = paragraph._p.getparent()
                if parent is not None:
                    parent.remove(paragraph._p)
        return len(elements)

    @staticmethod
    def _replace_in_paragraph(
        paragraph: Any,
        old: str,
        new: str,
        *,
        replace_all: bool,
    ) -> int:
        """Replace text across runs while retaining unaffected run formatting."""
        runs = list(paragraph.runs)
        if not runs:
            return 0
        text = "".join(run.text for run in runs)
        positions: list[int] = []
        cursor = 0
        while True:
            position = text.find(old, cursor)
            if position < 0:
                break
            positions.append(position)
            if not replace_all:
                break
            cursor = position + len(old)
        if not positions:
            return 0

        for start in reversed(positions):
            end = start + len(old)
            boundaries: list[tuple[int, int]] = []
            offset = 0
            for run in runs:
                next_offset = offset + len(run.text)
                boundaries.append((offset, next_offset))
                offset = next_offset

            start_index = next(
                index
                for index, (_, run_end) in enumerate(boundaries)
                if start < run_end
            )
            end_index = next(
                index
                for index, (_, run_end) in enumerate(boundaries)
                if end <= run_end
            )
            start_offset = start - boundaries[start_index][0]
            end_offset = end - boundaries[end_index][0]
            prefix = runs[start_index].text[:start_offset]
            suffix = runs[end_index].text[end_offset:]

            if start_index == end_index:
                runs[start_index].text = prefix + new + suffix
                continue
            runs[start_index].text = prefix + new
            for index in range(start_index + 1, end_index):
                runs[index].text = ""
            runs[end_index].text = suffix
        return len(positions)

    def _replace_text(
        self,
        document: Any,
        old: str,
        new: str,
        *,
        replace_all: bool,
        required: bool = True,
    ) -> int:
        count = 0
        for paragraph in self._paragraphs(document):
            replacements = self._replace_in_paragraph(
                paragraph,
                old,
                new,
                replace_all=replace_all,
            )
            count += replacements
            if replacements and not replace_all:
                break
        if count == 0 and required:
            raise ValueError(f"Word 文档中没有找到要替换的文字: {old}")
        return count
